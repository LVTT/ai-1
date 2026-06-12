"""文档加载模块

支持多种文档格式的加载与文本提取。
"""

import os
from pathlib import Path
from typing import List, Optional, Union
from dataclasses import dataclass


@dataclass
class Document:
    """统一文档数据结构"""
    content: str
    metadata: dict
    source: str


class DocumentLoader:
    """文档加载器

    支持格式：txt, md, pdf, docx
    """

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    @staticmethod
    def _extract_metadata(text: str) -> dict:
        """从原始文本中提取需求追溯关键信息"""
        import re

        # 提取 @人名
        people = list(set(re.findall(r'@([\u4e00-\u9fff\w]+)', text)))

        # 提取 【版本/分支】
        versions = list(set(re.findall(r'【([^】]+)】', text)))

        # 提取 [状态标签]
        statuses = list(set(re.findall(r'\[([^\]]+)\]', text)))

        # 提取日期（如 0829、2024-03-15）
        dates = list(set(re.findall(
            r'\b(?:20\d{2}[\-/])?\d{1,2}[\-/]\d{1,2}\b|\b\d{4}[\-/]\d{2}[\-/]\d{2}\b', text)))

        meta = {}
        if people:
            meta["people"] = ", ".join(people)
        if versions:
            meta["versions"] = ", ".join(versions)
        if statuses:
            meta["statuses"] = ", ".join(statuses)
        if dates:
            meta["dates"] = ", ".join(dates)

        return meta

    @staticmethod
    def _clean_text(text: str) -> str:
        """清洗文本噪音

        去除需求文档中常见的干扰信息：
        - @人名 提及
        - 【开发分支】等标记
        - [一期已上线] 等状态标签
        - 纯日期/数字行
        - 多余空行
        """
        import re

        lines = text.split("\n")
        cleaned_lines = []

        for line in lines:
            original = line.strip()
            if not original:
                continue  # 跳过空行

            # 去除 @人名（如 @钱金玉、@杜广宁@仲逸明）
            line = re.sub(r'@[\u4e00-\u9fff\w]+', '', line)

            # 去除 【...】内容
            line = re.sub(r'【[^】]*】', '', line)

            # 去除 [...] 标签（如 [一期已上线]）
            line = re.sub(r'\[[^\]]*\]', '', line)

            # 去除日期前缀行（如 0829：、0819~0823）
            if re.match(r'^[\d]{2,4}[~\-]?[\d]{0,4}\s*[:：]?\s*$', line.strip()):
                continue

            # 去除纯数字行
            if re.match(r'^[\d\s~\-:]+$', line.strip()):
                continue

            # 去除多余空格
            line = re.sub(r'\s+', ' ', line).strip()

            if line:
                cleaned_lines.append(line)

        return "\n".join(cleaned_lines)

    def load(self, source: Union[str, Path]) -> List[Document]:
        """加载单个文件或目录"""
        source = Path(source)

        if source.is_dir():
            return self.load_directory(source)
        else:
            doc = self.load_file(source)
            return [doc] if doc else []

    def load_file(self, file_path: Union[str, Path]) -> Optional[Document]:
        """加载单个文件"""
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()

        try:
            if suffix == ".txt" or suffix == ".md":
                return self._load_text(file_path)
            elif suffix == ".pdf":
                return self._load_pdf(file_path)
            elif suffix == ".docx":
                return self._load_docx(file_path)
            else:
                print(f"不支持的文件格式: {suffix}")
                return None
        except Exception as e:
            print(f"加载文件失败 {file_path}: {e}")
            return None

    def load_directory(self, dir_path: Union[str, Path]) -> List[Document]:
        """加载目录下所有支持的文件"""
        dir_path = Path(dir_path)
        documents = []

        for file_path in dir_path.rglob("*"):
            if file_path.is_file() and file_path.suffix.lower() in [".txt", ".md", ".pdf", ".docx"]:
                doc = self.load_file(file_path)
                if doc:
                    documents.append(doc)

        return documents

    def _load_text(self, file_path: Path) -> Document:
        """加载文本文件"""
        with open(file_path, "r", encoding="utf-8") as f:
            raw_content = f.read()

        extra_meta = self._extract_metadata(raw_content)
        content = self._clean_text(raw_content)

        return Document(
            content=content,
            metadata={
                "file_name": file_path.name,
                "file_type": "text",
                "file_path": str(file_path),
                **extra_meta,
            },
            source=str(file_path),
        )

    def _load_pdf(self, file_path: Path) -> Document:
        """加载 PDF 文件"""
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("请安装 pypdf: pip install pypdf")

        reader = PdfReader(str(file_path))
        texts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                texts.append(text)

        content = "\n".join(texts)
        return Document(
            content=content,
            metadata={
                "file_name": file_path.name,
                "file_type": "pdf",
                "file_path": str(file_path),
                "page_count": len(reader.pages),
            },
            source=str(file_path),
        )

    def _load_docx(self, file_path: Path) -> Document:
        """加载 Word 文档"""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")

        doc = DocxDocument(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        raw_content = "\n".join(paragraphs)

        extra_meta = self._extract_metadata(raw_content)
        content = self._clean_text(raw_content)

        # 从文件名提取年份和模块
        import re
        year_match = re.search(r'(20\d{2})', file_path.name)
        module_match = re.search(r'(CRM|销售系统|办公系统|费用管理|报价)', file_path.name)

        return Document(
            content=content,
            metadata={
                "file_name": file_path.name,
                "file_type": "docx",
                "file_path": str(file_path),
                "year": year_match.group(1) if year_match else None,
                "module": module_match.group(1) if module_match else None,
                **extra_meta,
            },
            source=str(file_path),
        )

    def split_text(self, text: str) -> List[str]:
        """简单文本分块（按字符数）

        生产环境建议使用更智能的分块策略，如：
        - LangChain 的 RecursiveCharacterTextSplitter
        - 按语义分块
        - 按句子/段落分块
        """
        chunks = []
        start = 0
        text_len = len(text)

        while start < text_len:
            end = start + self.chunk_size
            chunk = text[start:end]
            chunks.append(chunk)
            start += self.chunk_size - self.chunk_overlap

        return chunks

    def load_and_split(self, source: Union[str, Path]) -> List[Document]:
        """加载文档并分块"""
        documents = self.load(source)
        split_docs = []

        for doc in documents:
            chunks = self.split_text(doc.content)
            total_len = len(doc.content)
            for i, chunk in enumerate(chunks):
                metadata = doc.metadata.copy()
                metadata["chunk_index"] = i
                metadata["total_chunks"] = len(chunks)
                # 计算这块在原文中的大致位置百分比
                chunk_start = i * (self.chunk_size - self.chunk_overlap)
                metadata["position_percent"] = round(
                    min(100, (chunk_start / max(total_len, 1)) * 100), 1
                )
                split_docs.append(Document(
                    content=chunk,
                    metadata=metadata,
                    source=doc.source,
                ))

        return split_docs
